"""Document loading and normalization for multi-format input files."""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import tempfile
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Tuple

import chardet  # type: ignore

SUPPORTED_INPUT_EXTS = {".txt", ".pdf", ".epub", ".docx", ".mobi", ".azw3"}
_HTML_EXTS = {".html", ".htm", ".xhtml"}


@dataclass
class PreparedDocument:
    source_path: str
    source_ext: str
    working_text_path: str
    encoding: str = "utf-8"
    native_chapters: Optional[List[Dict[str, Any]]] = None
    has_native_structure: bool = False
    temp_dirs: List[str] = field(default_factory=list)


class DocumentLoader:
    def __init__(self) -> None:
        self._cache_root = os.path.join(tempfile.gettempdir(), "txt_splitter_cache")
        os.makedirs(self._cache_root, exist_ok=True)
        self._managed_files: set[str] = set()

    def prepare(self, file_path: str) -> PreparedDocument:
        abs_path = os.path.abspath(file_path)
        ext = os.path.splitext(abs_path)[1].lower()
        if ext not in SUPPORTED_INPUT_EXTS:
            raise ValueError(f"Unsupported input format: {ext}")
        if not os.path.exists(abs_path):
            raise FileNotFoundError(abs_path)

        if ext == ".txt":
            enc = self.detect_encoding(abs_path)
            return PreparedDocument(
                source_path=abs_path,
                source_ext=ext,
                working_text_path=abs_path,
                encoding=enc,
                native_chapters=None,
                has_native_structure=False,
            )

        if ext == ".pdf":
            text = self._extract_pdf_text(abs_path)
            return self._build_prepared(abs_path, ext, text, None, [])

        if ext == ".docx":
            text, chapters = self._extract_docx_text_and_headings(abs_path)
            return self._build_prepared(abs_path, ext, text, chapters, [])

        if ext == ".epub":
            text, chapters = self._extract_epub_text_and_toc(abs_path)
            return self._build_prepared(abs_path, ext, text, chapters, [])

        if ext in {".mobi", ".azw3"}:
            return self._prepare_mobi_like(abs_path, ext)

        raise ValueError(f"Unsupported input format: {ext}")

    def cleanup(self, doc: PreparedDocument) -> None:
        working_path = os.path.abspath(doc.working_text_path)
        if working_path in self._managed_files:
            self._managed_files.discard(working_path)
            try:
                if os.path.exists(working_path):
                    os.remove(working_path)
            except OSError:
                pass

        for temp_dir in doc.temp_dirs:
            try:
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            except OSError:
                pass

    def cleanup_all(self) -> None:
        for path in list(self._managed_files):
            try:
                if os.path.exists(path):
                    os.remove(path)
            except OSError:
                pass
            self._managed_files.discard(path)

    def detect_encoding(self, file_path: str) -> str:
        with open(file_path, "rb") as f:
            raw_data = f.read(10000)
        result = chardet.detect(raw_data)
        encoding = result.get("encoding")
        if not encoding or result.get("confidence", 0.0) < 0.7:
            return "utf-8"
        if encoding.lower() in {"gb2312", "gb18030"}:
            return "gbk"
        return encoding

    def _build_prepared(
        self,
        source_path: str,
        source_ext: str,
        text: str,
        native_chapters: Optional[List[Dict[str, Any]]],
        temp_dirs: List[str],
    ) -> PreparedDocument:
        normalized = self._normalize_text(text)
        if not normalized.strip():
            raise ValueError(f"No text content extracted from {source_ext} file.")
        working = self._write_working_text(source_path, normalized)
        valid_native = native_chapters if native_chapters and len(native_chapters) >= 2 else None
        return PreparedDocument(
            source_path=source_path,
            source_ext=source_ext,
            working_text_path=working,
            encoding="utf-8",
            native_chapters=valid_native,
            has_native_structure=valid_native is not None,
            temp_dirs=temp_dirs,
        )

    def _write_working_text(self, source_path: str, text: str) -> str:
        out_path = self._working_path_for(source_path)
        with open(out_path, "w", encoding="utf-8", newline="\n") as f:
            f.write(text)
        self._managed_files.add(out_path)
        return out_path

    def _working_path_for(self, source_path: str) -> str:
        st = os.stat(source_path)
        key = f"{os.path.abspath(source_path)}|{st.st_mtime_ns}|{st.st_size}"
        digest = hashlib.sha1(key.encode("utf-8")).hexdigest()[:12]
        base = os.path.splitext(os.path.basename(source_path))[0]
        safe_base = re.sub(r"[^0-9A-Za-z._-]+", "_", base).strip("_") or "document"
        return os.path.join(self._cache_root, f"{safe_base}_{digest}.txt")

    def _normalize_text(self, text: str) -> str:
        cleaned = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
        cleaned = "\n".join(line.rstrip() for line in cleaned.split("\n"))
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = cleaned.strip()
        return cleaned + ("\n" if cleaned else "")

    def _extract_pdf_text(self, file_path: str) -> str:
        probe_pages = 30
        pdfplumber_error: Optional[Exception] = None
        try:
            import pdfplumber  # type: ignore

            pages: List[str] = []
            empty_count = 0
            with pdfplumber.open(file_path) as pdf:
                for idx, page in enumerate(pdf.pages):
                    text = page.extract_text(layout=True) or page.extract_text() or ""
                    if text.strip():
                        pages.append(text)
                    else:
                        empty_count += 1
                    # Fast-fail for scanned/image PDFs with no text layer.
                    if (idx + 1) >= probe_pages and not pages:
                        raise ValueError(f"pdfplumber probe found no text in first {probe_pages} pages.")
            joined = "\n\n".join(pages)
            if not joined.strip():
                raise ValueError("pdfplumber returned no text.")
            if pages and empty_count / max(len(pages) + empty_count, 1) > 0.85:
                raise ValueError("pdfplumber text density is too low.")
            return joined
        except Exception as exc:  # noqa: BLE001
            pdfplumber_error = exc

        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(file_path)
            pages: List[str] = []
            nonempty_seen = False
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages.append(text)
                if text.strip():
                    nonempty_seen = True
                if (idx + 1) >= probe_pages and not nonempty_seen:
                    raise ValueError(f"pypdf probe found no text in first {probe_pages} pages.")

            joined = "\n\n".join(pages)
            if not joined.strip():
                raise ValueError("pypdf returned no text.")
            return joined
        except Exception as exc:  # noqa: BLE001
            pypdf_error = exc
            # Last fallback: pdfminer can sometimes extract text from PDFs that
            # pdfplumber/pypdf cannot.
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract_text  # type: ignore

                text = pdfminer_extract_text(file_path) or ""
                text = text.replace("\x0c", "\n").strip()
                if self._has_meaningful_text(text):
                    return text
                raise ValueError("pdfminer returned no meaningful text.")
            except Exception as pdfminer_exc:  # noqa: BLE001
                if pdfplumber_error:
                    raise ValueError(
                        "No extractable text found in PDF. "
                        "Tried pdfplumber, pypdf, and pdfminer.\n"
                        f"pdfplumber: {pdfplumber_error}\n"
                        f"pypdf: {pypdf_error}\n"
                        f"pdfminer: {pdfminer_exc}"
                    ) from pdfminer_exc
                raise

    def _has_meaningful_text(self, text: str) -> bool:
        if not text:
            return False
        # Ignore whitespace/control characters and require a minimal amount of content.
        compact = re.sub(r"[\s\x00-\x1f]+", "", text)
        return len(compact) >= 40

    def _extract_docx_text_and_headings(
        self, file_path: str
    ) -> Tuple[str, Optional[List[Dict[str, Any]]]]:
        import docx  # type: ignore

        document = docx.Document(file_path)
        lines: List[str] = []
        chapter_entries: List[Dict[str, Any]] = []
        level_stack: List[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            line_start = len(lines)
            lines.append(text)

            style_name = ""
            if para.style is not None and para.style.name:
                style_name = str(para.style.name)
            level = self._heading_level(style_name)
            if not level:
                continue

            safe_title = self._sanitize_title(text)
            level_stack = level_stack[: level - 1]
            level_stack.append(safe_title)
            title_prefix = "".join(f"[{p}] " for p in level_stack[:-1])
            chapter_entries.append(
                {
                    "title": f"{title_prefix}{safe_title}".strip(),
                    "raw_title": safe_title,
                    "hierarchy_path": list(level_stack),
                    "line_start": line_start,
                    "line_end": -1,
                }
            )

        chapters = self._finalize_chapter_ranges(chapter_entries, len(lines))
        return ("\n".join(lines)), chapters

    def _heading_level(self, style_name: str) -> Optional[int]:
        if not style_name:
            return None
        m = re.search(r"(?:heading|标题)\s*([1-6])", style_name, re.IGNORECASE)
        if not m:
            return None
        level = int(m.group(1))
        return level if 1 <= level <= 6 else None

    def _extract_epub_text_and_toc(
        self, file_path: str
    ) -> Tuple[str, Optional[List[Dict[str, Any]]]]:
        import ebooklib  # type: ignore
        from ebooklib import epub  # type: ignore

        book = epub.read_epub(file_path)
        lines: List[str] = []
        href_to_line: Dict[str, int] = {}

        for spine_item in book.spine:
            item_id = spine_item[0] if isinstance(spine_item, (tuple, list)) else spine_item
            item = book.get_item_with_id(item_id)
            if not item or item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue

            href = self._normalize_href(getattr(item, "file_name", "") or "")
            if href:
                href_to_line[href] = len(lines)

            raw_content = item.get_content()
            html = raw_content.decode("utf-8", errors="ignore") if isinstance(raw_content, bytes) else str(raw_content)
            text = self._extract_html_text(html)
            if text.strip():
                lines.extend(text.splitlines())
                lines.append("")

        chapters = self._epub_toc_to_chapters(getattr(book, "toc", []), href_to_line, len(lines))
        return ("\n".join(lines)), chapters

    def _epub_toc_to_chapters(
        self, toc: Any, href_to_line: Dict[str, int], total_lines: int
    ) -> Optional[List[Dict[str, Any]]]:
        flat: List[Tuple[int, str, str]] = []

        def walk(nodes: Sequence[Any], level: int) -> None:
            for node in nodes:
                if isinstance(node, tuple) and len(node) == 2 and isinstance(node[1], (list, tuple)):
                    title, href = self._toc_title_href(node[0])
                    if title:
                        flat.append((level, title, href))
                    walk(node[1], level + 1)
                    continue

                title, href = self._toc_title_href(node)
                if title:
                    flat.append((level, title, href))

        if isinstance(toc, (list, tuple)):
            walk(toc, 1)
        elif toc:
            walk([toc], 1)

        entries: List[Dict[str, Any]] = []
        stack: List[str] = []
        for level, title, href in flat:
            line_start = href_to_line.get(self._normalize_href(href))
            if line_start is None:
                continue
            safe = self._sanitize_title(title)
            if not safe:
                continue
            stack = stack[: max(level - 1, 0)]
            stack.append(safe)
            title_prefix = "".join(f"[{p}] " for p in stack[:-1])
            entries.append(
                {
                    "title": f"{title_prefix}{safe}".strip(),
                    "raw_title": safe,
                    "hierarchy_path": list(stack),
                    "line_start": line_start,
                    "line_end": -1,
                }
            )

        return self._finalize_chapter_ranges(entries, total_lines)

    def _toc_title_href(self, node: Any) -> Tuple[str, str]:
        title = ""
        href = ""
        if node is None:
            return title, href
        if hasattr(node, "title"):
            title = str(getattr(node, "title") or "")
        if hasattr(node, "href"):
            href = str(getattr(node, "href") or "")
        if not title and hasattr(node, "get_name"):
            try:
                title = str(node.get_name() or "")
            except Exception:  # noqa: BLE001
                title = ""
        if not href and hasattr(node, "file_name"):
            href = str(getattr(node, "file_name") or "")
        return title.strip(), href.strip()

    def _extract_html_text(self, html: str) -> str:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style"]):
            tag.extract()
        text = soup.get_text("\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text

    def _prepare_mobi_like(self, file_path: str, source_ext: str) -> PreparedDocument:
        import mobi  # type: ignore

        temp_dir, extracted_path = mobi.extract(file_path)
        temp_dirs = [temp_dir] if temp_dir else []
        target_path = self._pick_mobi_target(extracted_path, temp_dir)
        target_ext = os.path.splitext(target_path)[1].lower()

        chapters: Optional[List[Dict[str, Any]]] = None
        if target_ext == ".epub":
            text, chapters = self._extract_epub_text_and_toc(target_path)
        elif target_ext in _HTML_EXTS:
            with open(target_path, "r", encoding="utf-8", errors="ignore") as f:
                text = self._extract_html_text(f.read())
        elif target_ext == ".txt":
            enc = self.detect_encoding(target_path)
            with open(target_path, "r", encoding=enc, errors="replace") as f:
                text = f.read()
        elif target_ext == ".pdf":
            text = self._extract_pdf_text(target_path)
        elif target_ext == ".docx":
            text, chapters = self._extract_docx_text_and_headings(target_path)
        else:
            raise ValueError(
                f"Unsupported extracted content from {source_ext}: {os.path.basename(target_path)}"
            )

        return self._build_prepared(file_path, source_ext, text, chapters, temp_dirs)

    def _pick_mobi_target(self, extracted_path: str, temp_dir: str) -> str:
        if extracted_path and os.path.isfile(extracted_path):
            ext = os.path.splitext(extracted_path)[1].lower()
            if ext in SUPPORTED_INPUT_EXTS or ext in _HTML_EXTS:
                return extracted_path

        priority = [".epub", ".xhtml", ".html", ".htm", ".txt", ".pdf", ".docx"]
        candidates: Dict[str, List[str]] = {k: [] for k in priority}
        if temp_dir and os.path.isdir(temp_dir):
            for root, _, files in os.walk(temp_dir):
                for name in files:
                    ext = os.path.splitext(name)[1].lower()
                    if ext in candidates:
                        candidates[ext].append(os.path.join(root, name))

        for ext in priority:
            if candidates[ext]:
                candidates[ext].sort(key=lambda p: (len(p), -os.path.getsize(p)))
                return candidates[ext][0]

        raise ValueError("Unable to locate usable extracted content. File may be encrypted/DRM protected.")

    def _normalize_href(self, href: str) -> str:
        if not href:
            return ""
        clean = href.split("#", 1)[0].strip()
        return clean.replace("\\", "/")

    def _sanitize_title(self, text: str) -> str:
        safe = re.sub(r'[\\/*?:"<>|]', "", text).strip()
        return safe[:120]

    def _finalize_chapter_ranges(
        self, chapters: List[Dict[str, Any]], total_lines: int
    ) -> Optional[List[Dict[str, Any]]]:
        if len(chapters) < 2:
            return None
        for idx, chapter in enumerate(chapters):
            start = int(chapter.get("line_start", 0))
            if idx + 1 < len(chapters):
                next_start = int(chapters[idx + 1].get("line_start", start))
                end = max(start, next_start - 1)
            else:
                end = max(start, total_lines - 1)
            chapter["line_end"] = end
        return chapters
