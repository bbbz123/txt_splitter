from __future__ import annotations

from pathlib import Path

import docx
import pytest
from ebooklib import epub


def _escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_simple_pdf(path: Path, lines: list[str]) -> None:
    content_lines = ["BT", "/F1 18 Tf", "50 780 Td"]
    for idx, line in enumerate(lines):
        if idx:
            content_lines.append("0 -24 Td")
        content_lines.append(f"({_escape_pdf_text(line)}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", errors="replace")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream",
    ]

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{idx} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )

    path.write_bytes(bytes(output))


def _write_markdown(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "# Book Title",
                "",
                "Intro line one.",
                "Intro line two.",
                "",
                "## Chapter One",
                "Body one line one.",
                "Body one line two.",
                "",
                "## Chapter Two",
                "Body two line one.",
                "Body two line two.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def _write_docx(path: Path) -> None:
    document = docx.Document()
    document.add_heading("Book Title", level=1)
    document.add_paragraph("Intro line one.")
    document.add_paragraph("Intro line two.")
    document.add_heading("Chapter One", level=2)
    document.add_paragraph("Body one line one.")
    document.add_paragraph("Body one line two.")
    document.add_heading("Chapter Two", level=2)
    document.add_paragraph("Body two line one.")
    document.add_paragraph("Body two line two.")
    document.save(path)


def _write_epub(path: Path) -> None:
    book = epub.EpubBook()
    book.set_identifier("sample-book")
    book.set_title("Sample Book")
    book.set_language("en")

    intro = epub.EpubHtml(title="Book Title", file_name="intro.xhtml", lang="en")
    intro.content = "<h1>Book Title</h1><p>Intro line one.</p><p>Intro line two.</p>"

    chapter_one = epub.EpubHtml(title="Chapter One", file_name="chapter1.xhtml", lang="en")
    chapter_one.content = (
        "<h1>Chapter One</h1><p>Body one line one.</p><p>Body one line two.</p>"
    )

    chapter_two = epub.EpubHtml(title="Chapter Two", file_name="chapter2.xhtml", lang="en")
    chapter_two.content = (
        "<h1>Chapter Two</h1><p>Body two line one.</p><p>Body two line two.</p>"
    )

    book.add_item(intro)
    book.add_item(chapter_one)
    book.add_item(chapter_two)
    book.toc = (
        epub.Link("intro.xhtml", "Book Title", "intro"),
        (epub.Section("Main"), (chapter_one, chapter_two)),
    )
    book.spine = ["nav", intro, chapter_one, chapter_two]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub.write_epub(str(path), book)


@pytest.fixture
def sample_markdown_path(tmp_path: Path) -> Path:
    path = tmp_path / "sample.md"
    _write_markdown(path)
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    _write_docx(path)
    return path


@pytest.fixture
def sample_epub_path(tmp_path: Path) -> Path:
    path = tmp_path / "sample.epub"
    _write_epub(path)
    return path


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    _write_simple_pdf(
        path,
        [
            "Chapter 1",
            "Alpha line one.",
            "Alpha line two.",
            "Chapter 2",
            "Beta line one.",
            "Beta line two.",
        ],
    )
    return path
